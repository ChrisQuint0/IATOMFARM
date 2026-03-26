import csv
import io
import genanki
import random
from fpdf import FPDF
from docx import Document

def create_csv_export(cards):
    """Generates an Anki-compatible CSV string."""
    output = io.StringIO()
    writer = csv.writer(output, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
    for card in cards:
        writer.writerow([card['question'], card['answer']])
    return output.getvalue()

def create_docx_export(cards):
    """Generates a styled Word document."""
    doc = Document()
    doc.add_heading('IATOMFARM Study Reviewer', 0)
    
    for i, card in enumerate(cards, 1):
        # Question
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {card['question']}")
        run.bold = True
        
        # Answer
        doc.add_paragraph(f"   Answer: {card['answer']}")
        doc.add_paragraph() # Spacer
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def create_pdf_export(cards):
    """Generates a PDF reviewer using fpdf2."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", 'B', 20)
    pdf.cell(0, 20, 'IATOMFARM Reviewer', ln=True, align='C')
    pdf.ln(10)
    
    for i, card in enumerate(cards, 1):
        pdf.set_font("helvetica", 'B', 12)
        pdf.multi_cell(0, 8, f"{i}. {card['question']}")
        
        pdf.set_font("helvetica", '', 11)
        pdf.multi_cell(0, 8, f"   Answer: {card['answer']}")
        pdf.ln(4)
        
        # Add a light line between items
        pdf.set_draw_color(230, 230, 230)
        pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
        pdf.ln(4)
        
    return pdf.output()

def create_apkg_export(cards, deck_name="IATOMFARM Deck"):
    """Generates an Anki .apkg package."""
    # Unique IDs for model and deck
    model_id = random.randrange(1 << 30, 1 << 31)
    deck_id = random.randrange(1 << 30, 1 << 31)
    
    my_model = genanki.Model(
        model_id,
        'IATOMFARM Simple Model',
        fields=[
            {'name': 'Question'},
            {'name': 'Answer'},
        ],
        templates=[
            {
                'name': 'Card 1',
                'qfmt': '<div style="font-family: Arial; font-size: 20px; text-align: center; padding: 20px;">{{Question}}</div>',
                'afmt': '{{FrontSide}}<hr id="answer"><div style="font-family: Arial; font-size: 20px; text-align: center; color: #4F46E5; font-weight: bold; padding: 20px;">{{Answer}}</div>',
            },
        ])
    
    my_deck = genanki.Deck(deck_id, deck_name)
    
    for card in cards:
        my_note = genanki.Note(
            model=my_model,
            fields=[card['question'], card['answer']])
        my_deck.add_note(my_note)
    
    # Write to a file stream
    file_stream = io.BytesIO()
    # genanki.Package.write_to_file expects a filename or file-like object
    # Actually, we can use Package(my_deck).write_to_file(stream)
    package = genanki.Package(my_deck)
    package.write_to_file(file_stream)
    file_stream.seek(0)
    return file_stream
